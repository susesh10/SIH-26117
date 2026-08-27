from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime
import os
import re

def create_approval_note(content: str, output_folder: str = "outputs") -> str:
    """
    Creates a clean and professional Word document from the approval note content.
    """
    os.makedirs(output_folder, exist_ok=True)

    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # ---------- Helper function ----------
    def add_heading_custom(text, level=1):
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        return heading

    # ---------- Title ----------
    title = doc.add_heading("Inspection Observation / Approval Note", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)

    # ---------- Clean the content ----------
    # Remove common markdown symbols
    content = content.replace("**", "").replace("###", "").replace("##", "").replace("#", "")
    content = re.sub(r'\n{3,}', '\n\n', content)  # Remove excessive empty lines

    lines = content.strip().split("\n")

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Detect section headers
        if re.match(r'^\d+\.\s+(Background|Key Findings|Observations|Recommendations|Conclusion)', line, re.IGNORECASE):
            add_heading_custom(line, level=2)
        elif line.lower().startswith("title:") or line.lower().startswith("date:") or line.lower().startswith("reference:"):
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.bold = True
            run.font.size = Pt(11)
        else:
            # Normal paragraph
            para = doc.add_paragraph(line)
            para.paragraph_format.space_after = Pt(6)
            for run in para.runs:
                run.font.size = Pt(11)

    # ---------- Signature Section ----------
    doc.add_paragraph("\n")
    doc.add_paragraph("─" * 50)

    sig1 = doc.add_paragraph()
    sig1.add_run("Prepared by: ").bold = True
    sig1.add_run("_______________________________")

    sig2 = doc.add_paragraph()
    sig2.add_run("Designation: ").bold = True
    sig2.add_run("Inspection Engineer")

    doc.add_paragraph()

    sig3 = doc.add_paragraph()
    sig3.add_run("Approved by: ").bold = True
    sig3.add_run("_______________________________")

    sig4 = doc.add_paragraph()
    sig4.add_run("Designation: ").bold = True
    sig4.add_run("Head of Maintenance / Section Head")

    # ---------- Save ----------
    filename = f"Approval_Note_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(output_folder, filename)
    doc.save(filepath)

    return filepath